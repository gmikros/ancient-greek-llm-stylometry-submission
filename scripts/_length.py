# -*- coding: utf-8 -*-
"""Measure manuscript length by section group."""
import os, re
from docx import Document
from docx.oxml.ns import qn

ROOT = r"C:\Users\USER01\Dropbox\Workplace\D\George\PAPERS\Paroysiaseis\Cyprus 2025\ag-llm-stylometry"
doc = Document(os.path.join(ROOT, "Ancient_Greek_LLM_Stylometry_MERGED.docx"))

def wc(s):
    return len(s.split())

groups = {"front (title/abstract/keywords)": 0, "main text (1-8)": 0,
          "declarations": 0, "references": 0, "appendix": 0}
table_words = 0

# classify paragraphs by current section
section = "front (title/abstract/keywords)"
for p in doc.paragraphs:
    sn = p.style.name if p.style else ""
    t = p.text.strip()
    if sn == "Heading 1":
        h = t.lower()
        if h.startswith("1") and "introduction" in h:
            section = "main text (1-8)"
        elif h.startswith("declarations"):
            section = "declarations"
        elif h.startswith("references"):
            section = "references"
        elif h.startswith("appendix"):
            section = "appendix"
    groups[section] += wc(t)

# table text words
for tb in doc.tables:
    for row in tb.rows:
        for c in row.cells:
            table_words += wc(c.text)

n_img = len(doc.element.body.findall('.//' + qn('w:drawing')))
n_tbl = len(doc.tables)
n_ref = 0
inref = False
for p in doc.paragraphs:
    sn = p.style.name if p.style else ""
    if sn == "Heading 1" and p.text.strip().lower().startswith("references"):
        inref = True; continue
    if sn == "Heading 1" and inref:
        break
    if inref and p.text.strip():
        n_ref += 1

body_words = groups["main text (1-8)"]
total_text = sum(groups.values())
total_incl_tables = total_text + table_words

print("=== WORD COUNT BY SECTION ===")
for k, v in groups.items():
    print("  %-32s %6d" % (k, v))
print("  %-32s %6d" % ("table cell text", table_words))
print("  " + "-"*40)
print("  %-32s %6d" % ("MAIN TEXT (Intro..Conclusion)", body_words))
print("  %-32s %6d" % ("ALL paragraph text", total_text))
print("  %-32s %6d" % ("ALL incl. table cells", total_incl_tables))
print()
print("Figures:", n_img, "| Tables:", n_tbl, "| Reference entries:", n_ref)
print()
# rough typeset-page estimates (Springer single-column ~ 550-650 words/typeset page incl. spacing;
# double-spaced manuscript ~ 250-300 words/page)
wpp_typeset = 600
wpp_ms = 275
narrative = body_words + groups["declarations"]
print("Rough page estimates (narrative text %d words):" % narrative)
print("  ~%.0f typeset journal pages (text only, excl. figures/tables/refs)" % (narrative / wpp_typeset))
print("  ~%.0f double-spaced manuscript pages (text only)" % (narrative / wpp_ms))
