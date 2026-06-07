# -*- coding: utf-8 -*-
import os, re, unicodedata
from docx import Document
from docx.oxml.ns import qn

ROOT = r"C:\Users\USER01\Dropbox\Workplace\D\George\PAPERS\Paroysiaseis\Cyprus 2025\ag-llm-stylometry"
doc = Document(os.path.join(ROOT, "Ancient_Greek_LLM_Stylometry_MERGED.docx"))
import _manuscript_content as C  # noqa

def fold(s):
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn").lower()

# 1. reference ordering with diacritic folding
print("=== REFERENCE ORDER (diacritic-folded) ===")
refs = C.REFERENCES
keys = [fold(r.split("(")[0]).strip() for r in refs]
violations = [(refs[i].split('(')[0].strip(), refs[i+1].split('(')[0].strip())
              for i in range(len(keys)-1) if keys[i] > keys[i+1]]
print("Out-of-order adjacent pairs (folded):", len(violations))
for a, b in violations:
    print("   !", a, "  >  ", b)

# 2. true duplicate full-entry check
seen = {}
dups = []
for r in refs:
    if r in seen:
        dups.append(r)
    seen[r] = 1
print("Exact duplicate entries:", len(dups))

# 3. confirm the 4 known overlaps appear exactly once
for name in ["Haug DTT", "Qi P", "Singh P", "Riemenschneider F, Frank A (2023a)"]:
    c = sum(1 for r in refs if r.startswith(name))
    print("  count startswith %-38s = %d" % (repr(name), c))
print("  Riemenschneider 2023b present:", any(r.startswith("Riemenschneider F, Frank A (2023b)") for r in refs))

# 4. literal sentinel leftovers anywhere in the document?
full = "\n".join(p.text for p in doc.paragraphs)
print("\n=== SENTINEL / ARTIFACT CHECK ===")
print("Literal '[[I]]' or '[[/I]]' leftovers:", full.count("[[I]]") + full.count("[[/I]]"))
print("Literal '&' anywhere:", full.count(" & "))

# 5. italic run present for 'd' in abstract?
def italic_tokens(p):
    res = []
    for r in p.runs:
        if r.italic:
            res.append(r.text)
    return res
ab = next(p for p in doc.paragraphs if p.text.startswith("We release a chunk-level corpus"))
print("Abstract italic runs:", italic_tokens(ab))

# 6. dump first 600 chars of abstract and a discussion paragraph
print("\n=== ABSTRACT (raw) ===")
print(ab.text[:700])
disc = next(p for p in doc.paragraphs if p.text.startswith("The central result is one of magnitude"))
print("\n=== DISCUSSION 7.1 italic runs:", italic_tokens(disc))

# 7. confirm no leftover scope/title-divider text
for needle in ["Prepared for submission", "Scope.", "Methods (Sections 3", "Results (Section 6)",
               "References (methods and tools)", "References (literature review)",
               "should be merged into a single deduplicated"]:
    hit = any(needle in p.text for p in doc.paragraphs)
    print("Leftover %-45s : %s" % (repr(needle), hit))
