# -*- coding: utf-8 -*-
"""Validate the merged manuscript."""
import os, re, sys
from docx import Document
from docx.oxml.ns import qn

ROOT = r"C:\Users\USER01\Dropbox\Workplace\D\George\PAPERS\Paroysiaseis\Cyprus 2025\ag-llm-stylometry"
DOC = os.path.join(ROOT, "Ancient_Greek_LLM_Stylometry_MERGED.docx")
out = open(os.path.join(ROOT, "scripts", "_extracted", "MERGED_outline.txt"), "w", encoding="utf-8")

def w(s=""):
    out.write(s + "\n")

doc = Document(DOC)
body = doc.element.body

# ---- counts ----
n_img = len(body.findall('.//' + qn('w:drawing'))) + len(body.findall('.//' + qn('w:pict')))
n_tbl = len(doc.tables)
n_par = len(doc.paragraphs)
w("FILE: %s" % DOC)
w("Paragraphs: %d | Tables: %d | Images(drawings): %d" % (n_par, n_tbl, n_img))
w("")

# ---- heading outline + figure/table caption presence ----
w("=== HEADING OUTLINE ===")
for p in doc.paragraphs:
    sn = p.style.name if p.style is not None else ""
    t = p.text.strip()
    if not t:
        continue
    if sn in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        indent = {"Title": "", "Heading 1": "", "Heading 2": "   ", "Heading 3": "      "}[sn]
        w("%-9s | %s%s" % (sn, indent, t[:90]))
w("")

# ---- figure & table caption lines ----
w("=== FIGURE / TABLE CAPTIONS (first 100 chars) ===")
for p in doc.paragraphs:
    t = p.text.strip()
    if re.match(r"^(Figure|Table)\s+\d+", t):
        w("  " + t[:100])
w("")

# ---- citation style checks (body, excluding the References section) ----
w("=== CITATION STYLE CHECK (should find NO APA leftovers in body) ===")
# find the References heading index to exclude the bibliography
texts = [(p.style.name if p.style is not None else "", p.text) for p in doc.paragraphs]
ref_idx = None
for i, (sn, t) in enumerate(texts):
    if sn == "Heading 1" and t.strip() == "References":
        ref_idx = i
        break
apa_comma = []   # (Author, 1999) style
amp = []         # & in body
for i, (sn, t) in enumerate(texts):
    if ref_idx is not None and i >= ref_idx:
        break  # skip bibliography + appendix
    for m in re.finditer(r"\([A-Z][A-Za-zÀ-ÿ'’.\- ]+,\s+\d{4}[a-z]?", t):
        apa_comma.append(m.group(0))
    if " & " in t:
        amp.append(t[:60])
w("APA '(Author, YEAR' leftovers in body: %d" % len(apa_comma))
for s in apa_comma[:20]:
    w("   " + s)
w("'&' occurrences in body: %d" % len(amp))
for s in amp[:20]:
    w("   " + s)
w("")

# ---- sample converted Related Work citations ----
w("=== SAMPLE RELATED-WORK SENTENCES WITH CITATIONS ===")
shown = 0
for sn, t in texts:
    if "Burrows" in t or "Holmes" in t or "Juola" in t or "Mikros 2025" in t:
        w("  " + t[:240])
        shown += 1
    if shown >= 5:
        break
w("")

# ---- references: count, first, last, dup check ----
w("=== REFERENCES ===")
if ref_idx is not None:
    refs = []
    for i in range(ref_idx + 1, len(texts)):
        sn, t = texts[i]
        if sn == "Heading 1":  # next section (Appendix)
            break
        if t.strip():
            refs.append(t.strip())
    w("Reference entries: %d" % len(refs))
    w("First: " + refs[0][:80])
    w("Last:  " + refs[-1][:80])
    # duplicate first-author+year keys
    keys = {}
    for r in refs:
        m = re.match(r"^([A-Za-zÀ-ÿ’'\- ]+?)\s+(?:et al\s+)?\(?(\d{4}[a-z]?)", r)
        k = (m.group(1).strip(), m.group(2)) if m else (r[:20], "")
        keys.setdefault(k, 0)
        keys[k] += 1
    dups = {k: v for k, v in keys.items() if v > 1}
    w("Duplicate (author,year) keys: %s" % (dups if dups else "NONE"))
    # alphabetical order check
    first_tokens = [r.split("(")[0].strip().lower() for r in refs]
    ordered = all(first_tokens[i] <= first_tokens[i+1] for i in range(len(first_tokens)-1))
    w("Alphabetical order OK: %s" % ordered)
w("")

# ---- front matter check ----
w("=== FRONT MATTER ===")
joined = "\n".join(t for _, t in texts[:12])
w("Has Abstract heading: %s" % any(t.strip() == "Abstract" for _, t in texts))
w("Has Keywords: %s" % any(t.strip().startswith("Keywords") for _, t in texts))
w("Title (first non-empty): " + next(t for _, t in texts if t.strip())[:90])

# ---- line numbering ----
sectPr = body.find(qn('w:sectPr'))
w("Line numbering present: %s" % (sectPr is not None and sectPr.find(qn('w:lnNumType')) is not None))

out.close()
print("Wrote outline to scripts/_extracted/MERGED_outline.txt")
print("Images:", n_img, "Tables:", n_tbl, "Paragraphs:", n_par)
