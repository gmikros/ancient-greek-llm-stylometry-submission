# -*- coding: utf-8 -*-
"""Build the Literature Review (Section 2 / Related Work) Word document and a
companion reference-verification document, styled to match the existing
Methods-and-Results draft (Times New Roman 11pt body; Calibri bold dark-blue
headings; APA 7th references with hanging indent)."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\USER01\Dropbox\Workplace\D\George\PAPERS\Paroysiaseis\Cyprus 2025\ag-llm-stylometry"

# ---------------------------------------------------------------- style helpers
def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), "Times New Roman")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    specs = {
        "Heading 1": dict(size=15, color="1F3864", italic=False, before=17, after=8, border=True),
        "Heading 2": dict(size=12.5, color="1F3864", italic=False, before=13, after=6, border=False),
        "Heading 3": dict(size=11, color="2E4A7A", italic=True, before=9, after=4, border=False),
    }
    for name, s in specs.items():
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.bold = True
        st.font.italic = s["italic"]
        st.font.size = Pt(s["size"])
        st.font.color.rgb = RGBColor.from_string(s["color"])
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.append(rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), "Calibri")
        st.paragraph_format.space_before = Pt(s["before"])
        st.paragraph_format.space_after = Pt(s["after"])
        st.paragraph_format.keep_with_next = True

def add_bottom_border(paragraph, color="8EAADB", sz="6", space="4"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_italic_runs(paragraph, text):
    """Render *...* segments as italic runs."""
    parts = text.split("*")
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        run.italic = (i % 2 == 1)
    return paragraph

def body(doc, text, justify=True):
    p = doc.add_paragraph()
    add_italic_runs(p, text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def heading(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    if level == 1:
        add_bottom_border(p)
    return p

def reference(doc, text):
    p = doc.add_paragraph()
    add_italic_runs(p, text)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

# ----------------------------------------------------------------------- content
INTRO = (
    "This study sits at the convergence of five research strands: computational "
    "stylometry and authorship attribution, the computational philology of Ancient "
    "Greek, language resources and pretrained models for Ancient Greek, large "
    "language models for low-resource and historical languages, and the detection "
    "of machine-generated text. We review each in turn, drawing out the methods and "
    "findings on which the present resource and evaluation build, and the gap they "
    "are designed to fill: a documented, chunk-level corpus that pairs human "
    "Attic-orator prose with state-of-the-art large-language-model rewrites, and a "
    "statistically rigorous account of how closely, and under what conditions, the "
    "machine prose reproduces the style of human Ancient Greek."
)

S21 = [
    (
        "Stylometry — the quantitative study of linguistic style — has a long "
        "pedigree in the digital humanities, from early word-length and function-word "
        "studies to modern machine-learning pipelines (Holmes, 1998). Its central "
        "premise, that authors leave measurable and relatively stable stylistic "
        "fingerprints in features largely below conscious control, underpins both "
        "authorship attribution and the human-versus-machine comparison pursued here. "
        "The most influential distance-based method is Burrows’s (2002) Delta, which "
        "ranks candidate authors by the mean absolute difference of standardised "
        "high-frequency word frequencies; its geometric and probabilistic foundations "
        "were subsequently clarified by Argamon (2008), and the wider family of Delta "
        "measures was systematically evaluated by Evert et al. (2017)."
    ),
    (
        "Comprehensive surveys (Juola, 2006; Koppel et al., 2009; Stamatatos, 2009) and "
        "comparative evaluations of techniques (Grieve, 2007) map the broader feature "
        "space — function words, character and word n-grams, part-of-speech and "
        "syntactic patterns, and lexical-diversity indices — together with the "
        "classifiers built upon it, much of which is now packaged in widely used tools "
        "such as the stylo package for R (Eder et al., 2016). The present study inherits "
        "this tradition’s feature families and its reliance on standardised effect "
        "sizes and distance metrics, but redirects them from their usual target — "
        "discriminating among human authors — toward a different contrast, that "
        "between human and machine writers of the same language."
    ),
]

S22 = [
    (
        "Quantitative authorship study of Greek is itself decades old: Michaelson and "
        "Morton’s (1972) one-word test for Greek writers is an early instance of "
        "statistical authorship discrimination on the classical corpus. Modern work has "
        "moved from word-frequency methods to syntactic stylometry over dependency "
        "treebanks, an approach especially attractive for a heavily inflected, "
        "relatively free-word-order language: Gorman and Gorman (2016) use syntactic "
        "n-grams to probe text reuse in Greek historiography, and Gorman (2020) shows "
        "that author identification of short Greek texts is feasible from "
        "dependency-treebank features alone, without recourse to vocabulary."
    ),
    (
        "Stylometric and computational methods have likewise been brought to bear on "
        "canonical authorship problems, including the Homeric question (Pavlopoulos & "
        "Konstantinidou, 2023) and the disputed Aeschylean authorship of Prometheus "
        "Bound (Manousakis, 2020). The same verification machinery has resolved "
        "high-profile cases in the sister classical language, attributing a newly "
        "discovered work to a major second-century author (Stover et al., 2016) and "
        "authenticating the writings transmitted under Caesar’s name (Kestemont et al., "
        "2016). Closer to the present author’s own work, stylometric attribution has "
        "been applied to Modern Greek social-media text (Mikros & Perifanos, 2013), and "
        "a tightly controlled bilingual study found authorship attribution to be "
        "markedly harder in Greek than in English — a gap attributed to Greek’s richer "
        "morphology (Juola et al., 2019). This last finding directly motivates two "
        "design choices in the present work: the prominence of morphological features "
        "(case, mood, tense and aspect) in our feature set, and the author-level random "
        "effects and author-grouped cross-validation used to keep the unevenly "
        "represented orators from dominating the estimates."
    ),
]

S23 = [
    (
        "The feasibility of a large-scale computational evaluation of Ancient Greek "
        "rests on more than a decade of resource building. Syntactically annotated "
        "treebanks — the Ancient Greek and Latin Dependency Treebanks (Bamman & Crane, "
        "2011), the PROIEL treebank (Haug & Jøhndal, 2008), and the Pedalion treebanks "
        "(Keersmaekers et al., 2019) — supply the training data for modern parsers, "
        "while large lemmatised corpora such as Diorisis (Vatri & McGillivray, 2018) and "
        "GLAUx (Keersmaekers, 2021) support distributional study. These annotations are "
        "now exposed through general toolkits, notably Stanza (Qi et al., 2020), whose "
        "Ancient Greek models are trained on the PROIEL treebank and provide the "
        "morphosyntactic analysis used here, and the Classical Language Toolkit "
        "(Johnson et al., 2021)."
    ),
    (
        "In parallel, transformer encoders pretrained on classical corpora have become "
        "available: Ancient-Greek-BERT (Singh et al., 2021), the GreBERTa and PhilBERTa "
        "models (Riemenschneider & Frank, 2023a), and, for the sister language, Latin "
        "BERT (Bamman & Burns, 2020); recent work continues to push morphological "
        "tagging and lemmatisation of Greek and Latin with such models (Celano, 2024; "
        "Keersmaekers & Mercelis, 2024). Our evaluation uses this infrastructure "
        "directly — Stanza’s PROIEL pipeline for the stylometric features and "
        "Ancient-Greek-BERT and GreBERTa for the neural document embeddings — and, as a "
        "resource paper for Language Resources and Evaluation, contributes back to the "
        "ecosystem a documented parallel corpus of human and machine-generated Attic "
        "prose."
    ),
]

S24 = [
    (
        "The systems evaluated here are general-purpose large language models whose "
        "few-shot and instruction-following abilities (Bommasani et al., 2021; Brown et "
        "al., 2020) make it possible to elicit Ancient Greek prose through prompting "
        "alone. Because the output is steered entirely by the prompt, the way an "
        "instruction is framed materially shapes what is generated (Wei et al., 2022), "
        "which motivates our explicit contrast between a Restricted (close-rewrite) and "
        "a Free prompting regime."
    ),
    (
        "Although multilingual models still perform unevenly on low-resource languages "
        "(Zhu et al., 2024), a growing body of work probes their competence on "
        "historical and classical languages specifically: cross-lingual models that "
        "link Latin to Ancient Greek (Riemenschneider & Frank, 2023b), the use of "
        "ChatGPT for endangered and low-resource languages (Hämäläinen, 2024), and "
        "expert evaluation of large-language-model translations of technical Ancient "
        "Greek such as the medical and philosophical prose of Galen (Zainaldin et al., "
        "2026). Most relevant to the question of stylistic fidelity, Mikros (2025a) "
        "shows that although GPT-4o can approximate the surface style of individual "
        "literary authors, its imitations remain stylometrically separable from genuine "
        "human writing. Prior work, however, concentrates on understanding or "
        "translating historical text, or on imitating modern authors; the complementary "
        "question of how faithfully state-of-the-art models can generate prose in the "
        "style of a specific historical register — measured stylometrically and at "
        "scale — is largely open, and is the question this resource is built to answer."
    ),
]

S25 = [
    (
        "Finally, the detectability analyses in this paper connect to the fast-moving "
        "literature on distinguishing machine-generated from human text. Early detectors "
        "exploited the statistical signature of model sampling, from the GLTR "
        "visualisation tool (Gehrmann et al., 2019) and the fine-tuned neural "
        "classifiers released with GPT-2 (Solaiman et al., 2019) to the observation that "
        "the most fluent generations are paradoxically the hardest for humans — though "
        "not always for machines — to flag (Ippolito et al., 2020). The field has since "
        "been mapped by several surveys (Crothers et al., 2023; Jawahar et al., 2020; "
        "Tang et al., 2024) and advanced by zero-shot, curvature-based methods such as "
        "DetectGPT (Mitchell et al., 2023), purpose-built comparison corpora such as HC3 "
        "(Guo et al., 2023), small-model detectors (Mireshghallah et al., 2024) and "
        "proactive watermarking (Kirchenbauer et al., 2023); at the same time, "
        "robustness studies caution that detection can degrade sharply under paraphrase "
        "and adversarial pressure (Sadasivan et al., 2023)."
    ),
    (
        "Stylometric features remain competitive and complementary in this setting, both "
        "as standalone signals and in ensembles with transformer representations (Mikros "
        "et al., 2023), with direct implications for forensic linguistics as generative "
        "models mature (Mikros, 2025b). The overwhelming majority of this work targets "
        "modern, high-resource English. The present study repurposes a detectability "
        "metric — the cross-validated area under the ROC curve of a human-versus-AI "
        "classifier — not as an end in itself but as a graded, interpretable measure of "
        "stylistic fidelity for a historical language, and shows that it varies "
        "systematically with model and prompting regime, thereby tying the detection "
        "question back to the stylometric one."
    ),
]

CLOSING = (
    "Across these five strands, no prior study, to our knowledge, evaluates how closely "
    "contemporary large language models reproduce the style of a historical language at "
    "corpus scale, with a nested statistical design that respects the structure of the "
    "data and an accompanying released resource. The remainder of the paper addresses "
    "this gap: Sections 3 to 5 describe the corpus, the features and the evaluation "
    "methodology, and Section 6 reports the results."
)

REFERENCES = [
    "Argamon, S. (2008). Interpreting Burrows’s Delta: Geometric and probabilistic foundations. *Literary and Linguistic Computing, 23*(2), 131–147. https://doi.org/10.1093/llc/fqn003",
    "Bamman, D., & Burns, P. J. (2020). *Latin BERT: A contextual language model for classical philology* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2009.10053",
    "Bamman, D., & Crane, G. (2011). The Ancient Greek and Latin dependency treebanks. In C. Sporleder, A. van den Bosch, & K. Zervanou (Eds.), *Language technology for cultural heritage* (pp. 79–98). Springer. https://doi.org/10.1007/978-3-642-20227-8_5",
    "Bommasani, R., et al. (2021). *On the opportunities and risks of foundation models* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2108.07258",
    "Brown, T. B., et al. (2020). Language models are few-shot learners. In *Advances in Neural Information Processing Systems 33* (pp. 1877–1901). Curran Associates.",
    "Burrows, J. (2002). ‘Delta’: A measure of stylistic difference and a guide to likely authorship. *Literary and Linguistic Computing, 17*(3), 267–287. https://doi.org/10.1093/llc/17.3.267",
    "Celano, G. G. A. (2024). *A state-of-the-art morphosyntactic parser and lemmatizer for Ancient Greek* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2410.12055",
    "Crothers, E. N., Japkowicz, N., & Viktor, H. L. (2023). Machine-generated text: A comprehensive survey of threat models and detection methods. *IEEE Access, 11*, 70977–71002. https://doi.org/10.1109/ACCESS.2023.3294090",
    "Eder, M., Rybicki, J., & Kestemont, M. (2016). Stylometry with R: A package for computational text analysis. *The R Journal, 8*(1), 107–121. https://doi.org/10.32614/RJ-2016-007",
    "Evert, S., Proisl, T., Jannidis, F., Reger, I., Pielström, S., Schöch, C., & Vitt, T. (2017). Understanding and explaining Delta measures for authorship attribution. *Digital Scholarship in the Humanities, 32*(Suppl. 2), ii4–ii16. https://doi.org/10.1093/llc/fqx023",
    "Gehrmann, S., Strobelt, H., & Rush, A. M. (2019). GLTR: Statistical detection and visualization of generated text. In *Proceedings of the 57th Annual Meeting of the Association for Computational Linguistics: System Demonstrations* (pp. 111–116). Association for Computational Linguistics. https://doi.org/10.18653/v1/P19-3019",
    "Gorman, R. (2020). Author identification of short texts using dependency treebanks without vocabulary. *Digital Scholarship in the Humanities, 35*(4), 812–825. https://doi.org/10.1093/llc/fqz070",
    "Gorman, V. B., & Gorman, R. J. (2016). Approaching questions of text reuse in Ancient Greek using computational syntactic stylometry. *Open Linguistics, 2*(1), 500–510. https://doi.org/10.1515/opli-2016-0026",
    "Grieve, J. (2007). Quantitative authorship attribution: An evaluation of techniques. *Literary and Linguistic Computing, 22*(3), 251–270. https://doi.org/10.1093/llc/fqm020",
    "Guo, B., et al. (2023). *How close is ChatGPT to human experts? Comparison corpus, evaluation, and detection* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2301.07597",
    "Hämäläinen, M. (2024). DAG: Dictionary-augmented generation for disambiguation of sentences in endangered Uralic languages using ChatGPT. In *Proceedings of the 9th International Workshop on Computational Linguistics for Uralic Languages* (pp. 36–40). Association for Computational Linguistics.",
    "Haug, D. T. T., & Jøhndal, M. L. (2008). Creating a parallel treebank of the old Indo-European Bible translations. In C. Sporleder & K. Ribarov (Eds.), *Proceedings of the Second Workshop on Language Technology for Cultural Heritage Data (LaTeCH 2008)* (pp. 27–34). European Language Resources Association.",
    "Holmes, D. I. (1998). The evolution of stylometry in humanities scholarship. *Literary and Linguistic Computing, 13*(3), 111–117. https://doi.org/10.1093/llc/13.3.111",
    "Ippolito, D., Duckworth, D., Callison-Burch, C., & Eck, D. (2020). Automatic detection of generated text is easiest when humans are fooled. In *Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics* (pp. 1808–1822). Association for Computational Linguistics. https://doi.org/10.18653/v1/2020.acl-main.164",
    "Jawahar, G., Abdul-Mageed, M., & Lakshmanan, L. V. S. (2020). Automatic detection of machine generated text: A critical survey. In *Proceedings of the 28th International Conference on Computational Linguistics* (pp. 2296–2309). International Committee on Computational Linguistics. https://doi.org/10.18653/v1/2020.coling-main.208",
    "Johnson, K. P., Burns, P. J., Stewart, J., Cook, T., Besnier, C., & Mattingly, W. J. B. (2021). The Classical Language Toolkit: An NLP framework for pre-modern languages. In *Proceedings of the 59th Annual Meeting of the Association for Computational Linguistics and the 11th International Joint Conference on Natural Language Processing: System Demonstrations* (pp. 20–29). Association for Computational Linguistics. https://doi.org/10.18653/v1/2021.acl-demo.3",
    "Juola, P. (2006). Authorship attribution. *Foundations and Trends in Information Retrieval, 1*(3), 233–334. https://doi.org/10.1561/1500000005",
    "Juola, P., Mikros, G. K., & Vinsick, S. (2019). A comparative assessment of the difficulty of authorship attribution in Greek and in English. *Journal of the Association for Information Science and Technology, 70*(1), 61–70. https://doi.org/10.1002/asi.24073",
    "Keersmaekers, A. (2021). The GLAUx corpus: Methodological issues in designing a long-term, diverse, multi-layered corpus of Ancient Greek. In *Proceedings of the 2nd International Workshop on Computational Approaches to Historical Language Change 2021* (pp. 39–50). Association for Computational Linguistics. https://doi.org/10.18653/v1/2021.lchange-1.6",
    "Keersmaekers, A., & Mercelis, W. (2024). Adapting transformer models to morphological tagging of two highly inflectional languages: A case study on Ancient Greek and Latin. In *Proceedings of the 1st Workshop on Machine Learning for Ancient Languages (ML4AL 2024)* (pp. 165–176). Association for Computational Linguistics. https://doi.org/10.18653/v1/2024.ml4al-1.17",
    "Keersmaekers, A., Mercelis, W., Swaelens, C., & Van Hal, T. (2019). Creating, enriching and valorizing treebanks of Ancient Greek. In *Proceedings of the 18th International Workshop on Treebanks and Linguistic Theories (SyntaxFest 2019)* (pp. 109–117). Association for Computational Linguistics. https://doi.org/10.18653/v1/W19-7812",
    "Kestemont, M., Stover, J., Koppel, M., Karsdorp, F., & Daelemans, W. (2016). Authenticating the writings of Julius Caesar. *Expert Systems with Applications, 63*, 86–96. https://doi.org/10.1016/j.eswa.2016.06.029",
    "Kirchenbauer, J., Geiping, J., Wen, Y., Katz, J., Miers, I., & Goldstein, T. (2023). A watermark for large language models. In *Proceedings of the 40th International Conference on Machine Learning* (Vol. 202, pp. 17061–17084). PMLR.",
    "Koppel, M., Schler, J., & Argamon, S. (2009). Computational methods in authorship attribution. *Journal of the American Society for Information Science and Technology, 60*(1), 9–26. https://doi.org/10.1002/asi.20961",
    "Manousakis, N. (2020). *‘Prometheus Bound’: A separate authorial trace in the Aeschylean corpus* (Trends in Classics – Supplementary Volumes, Vol. 98). De Gruyter. https://doi.org/10.1515/9783110687675",
    "Michaelson, S., & Morton, A. Q. (1972). The new stylometry: A one-word test of authorship for Greek writers. *The Classical Quarterly, 22*(1), 89–102. https://doi.org/10.1017/S0009838800034054",
    "Mikros, G. K. (2025a). Beyond the surface: Stylometric analysis of GPT-4o’s capacity for literary style imitation. *Digital Scholarship in the Humanities, 40*(2), 587–600. https://doi.org/10.1093/llc/fqaf035",
    "Mikros, G. K. (2025b). *Large language models and forensic linguistics: Navigating opportunities and threats in the age of generative AI* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2512.06922",
    "Mikros, G. K., Koursaris, A., Bilianos, D., & Markopoulos, G. (2023). AI-writing detection using an ensemble of transformers and stylometric features. In *Proceedings of the Iberian Languages Evaluation Forum (IberLEF 2023)* (CEUR Workshop Proceedings, Vol. 3496). CEUR-WS. https://ceur-ws.org/Vol-3496/autextification-paper9.pdf",
    "Mikros, G. K., & Perifanos, K. (2013). Authorship attribution in Greek tweets using author’s multilevel n-gram profiles. In *Analyzing microtext: Papers from the 2013 AAAI Spring Symposium* (Technical Report SS-13-01, pp. 17–23). AAAI Press.",
    "Mireshghallah, N., Mattern, J., Gao, S., Shokri, R., & Berg-Kirkpatrick, T. (2024). Smaller language models are better zero-shot machine-generated text detectors. In *Proceedings of the 18th Conference of the European Chapter of the Association for Computational Linguistics (Volume 2: Short Papers)* (pp. 278–293). Association for Computational Linguistics. https://doi.org/10.18653/v1/2024.eacl-short.25",
    "Mitchell, E., Lee, Y., Khazatsky, A., Manning, C. D., & Finn, C. (2023). DetectGPT: Zero-shot machine-generated text detection using probability curvature. In *Proceedings of the 40th International Conference on Machine Learning* (Vol. 202, pp. 24950–24962). PMLR.",
    "Pavlopoulos, J., & Konstantinidou, M. (2023). Computational authorship analysis of the Homeric poems. *International Journal of Digital Humanities, 5*(1), 45–64. https://doi.org/10.1007/s42803-022-00046-7",
    "Qi, P., Zhang, Y., Zhang, Y., Bolton, J., & Manning, C. D. (2020). Stanza: A Python natural language processing toolkit for many human languages. In *Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics: System Demonstrations* (pp. 101–108). Association for Computational Linguistics. https://doi.org/10.18653/v1/2020.acl-demos.14",
    "Riemenschneider, F., & Frank, A. (2023a). Exploring large language models for classical philology. In *Proceedings of the 61st Annual Meeting of the Association for Computational Linguistics (Volume 1: Long Papers)* (pp. 15181–15199). Association for Computational Linguistics. https://doi.org/10.18653/v1/2023.acl-long.846",
    "Riemenschneider, F., & Frank, A. (2023b). Graecia capta ferum victorem cepit: Detecting Latin allusions to Ancient Greek literature. In *Proceedings of the Ancient Language Processing Workshop* (pp. 30–38). INCOMA.",
    "Sadasivan, V. S., Kumar, A., Balasubramanian, S., Wang, W., & Feizi, S. (2023). *Can AI-generated text be reliably detected?* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2303.11156",
    "Singh, P., Rutten, G., & Lefever, E. (2021). A pilot study for BERT language modelling and morphological analysis for ancient and medieval Greek. In *Proceedings of the 5th Joint SIGHUM Workshop on Computational Linguistics for Cultural Heritage, Social Sciences, Humanities and Literature* (pp. 128–137). Association for Computational Linguistics. https://doi.org/10.18653/v1/2021.latechclfl-1.15",
    "Solaiman, I., et al. (2019). *Release strategies and the social impacts of language models* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.1908.09203",
    "Stamatatos, E. (2009). A survey of modern authorship attribution methods. *Journal of the American Society for Information Science and Technology, 60*(3), 538–556. https://doi.org/10.1002/asi.21001",
    "Stover, J. A., Winter, Y., Koppel, M., & Kestemont, M. (2016). Computational authorship verification method attributes a new work to a major 2nd century African author. *Journal of the Association for Information Science and Technology, 67*(1), 239–242. https://doi.org/10.1002/asi.23460",
    "Tang, R., Chuang, Y.-N., & Hu, X. (2024). The science of detecting LLM-generated text. *Communications of the ACM, 67*(4), 50–59. https://doi.org/10.1145/3624725",
    "Vatri, A., & McGillivray, B. (2018). The Diorisis Ancient Greek Corpus. *Research Data Journal for the Humanities and Social Sciences, 3*(1), 55–65. https://doi.org/10.1163/24523666-01000013",
    "Wei, J., et al. (2022). Chain-of-thought prompting elicits reasoning in large language models. In *Advances in Neural Information Processing Systems 35* (pp. 24824–24837). Curran Associates.",
    "Zainaldin, J. L., Pattison, C., Marai, M., Wu, J., & Schiefsky, M. J. (2026). *Evaluating LLM-based translation of a low-resource technical language: The medical and philosophical Greek of Galen* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2602.24119",
    "Zhu, S., et al. (2024). *Multilingual large language models: A systematic survey* [Preprint]. arXiv. https://doi.org/10.48550/arXiv.2411.11072",
]

OVERLAP_NOTE = (
    "*Note.* Four entries above — Haug and Jøhndal (2008), Qi et al. (2020), Singh et "
    "al. (2021) and Riemenschneider and Frank (2023a) — also appear among the "
    "Methods-and-tools references in the existing draft and should be merged into a "
    "single deduplicated bibliography when the section is integrated. For works with "
    "eight or more authors, the abbreviated “first author et al.” form is used in the "
    "reference list to match the existing bibliography style (e.g., Pedregosa et al., "
    "2011); the full author lists are available from the cited DOIs and can be expanded "
    "to APA 7’s 20-author limit if the final venue requires it."
)

# ------------------------------------------------------------------ build doc 1
doc = Document()
configure_styles(doc)
section = doc.sections[0]
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Title block
t1 = doc.add_paragraph(style="Title")
r = t1.add_run("How Well Do State-of-the-Art LLMs Write Ancient Greek?")
t1.alignment = WD_ALIGN_PARAGRAPH.LEFT
t2 = doc.add_paragraph()
r = t2.add_run("A Chunk-Level Stylometric Resource and Evaluation")
r.bold = True; r.font.size = Pt(14); r.font.name = "Times New Roman"
auth = doc.add_paragraph()
auth.add_run("George Mikros").font.size = Pt(11)
ven = doc.add_paragraph()
rv = ven.add_run("Prepared for submission to Language Resources and Evaluation (Springer)")
rv.italic = True; rv.font.size = Pt(11)

scope = doc.add_paragraph()
sr = scope.add_run(
    "Scope. This document contains the Related Work / Literature Review (Section 2) for "
    "the manuscript, to be placed before the Methods (Sections 3–5). It is organised "
    "thematically; in-text citations and the reference list follow APA 7th edition, "
    "matching the rest of the paper.")
sr.italic = True
scope.paragraph_format.space_after = Pt(12)

heading(doc, "2  Related Work", 2)
body(doc, INTRO)

heading(doc, "2.1  Stylometry and authorship attribution", 3)
for para in S21:
    body(doc, para)

heading(doc, "2.2  Computational philology and stylometry of Ancient Greek", 3)
for para in S22:
    body(doc, para)

heading(doc, "2.3  Language resources and models for Ancient Greek", 3)
for para in S23:
    body(doc, para)

heading(doc, "2.4  Large language models for low-resource and historical languages", 3)
for para in S24:
    body(doc, para)

heading(doc, "2.5  Detecting and distinguishing machine-generated text", 3)
for para in S25:
    body(doc, para)

heading(doc, "2.6  Summary", 3)
body(doc, CLOSING)

heading(doc, "References (literature review)", 2)
for ref in REFERENCES:
    reference(doc, ref)

note_p = doc.add_paragraph()
add_italic_runs(note_p, OVERLAP_NOTE)
note_p.paragraph_format.space_before = Pt(8)
note_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

out1 = os.path.join(OUT_DIR, "Ancient_Greek_LLM_Stylometry_Literature_Review.docx")
doc.save(out1)
print("Saved:", out1)
print("Reference count:", len(REFERENCES))
