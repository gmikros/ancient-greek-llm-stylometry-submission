# -*- coding: utf-8 -*-
"""
Assemble the single, submission-ready LRE manuscript.

Strategy: start from a COPY of the Methods/Results .docx (so its 12 figures and
3 tables are preserved natively, with correct relationships), then:
  * strip its title block + scope + "Methods (Sections 3-5)" divider + bridge;
  * strip its "References (methods and tools)" section;
  * normalise heading levels to a flat 1-8 scheme;
  * insert front matter (title page, abstract, keywords) + Section 1 Introduction
    + Section 2 Related Work (cloned from the Literature Review .docx, with APA->
    Springer in-text citation conversion and italics preserved);
  * append Section 7 Discussion, Section 8 Conclusion, Declarations, the merged &
    deduplicated References (Springer author-date), and Appendix A (prompts);
  * add continuous line numbers for review.
Output: Ancient_Greek_LLM_Stylometry_MERGED.docx
"""
import os
import re
import shutil
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.text.paragraph import Paragraph

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import _manuscript_content as C

ROOT = r"C:\Users\USER01\Dropbox\Workplace\D\George\PAPERS\Paroysiaseis\Cyprus 2025\ag-llm-stylometry"
SRC_MR = os.path.join(ROOT, "Ancient_Greek_LLM_Stylometry_Methods_Results.docx")
SRC_LR = os.path.join(ROOT, "Ancient_Greek_LLM_Stylometry_Literature_Review.docx")
OUT = os.path.join(ROOT, "Ancient_Greek_LLM_Stylometry_MERGED.docx")
TMP = os.path.join(ROOT, "scripts", "_merge_tmp.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def ptext(p_el):
    return "".join(n.text or "" for n in p_el.iter(qn('w:t')))

def style_name(p):
    try:
        return p.style.name if p.style is not None else ""
    except Exception:
        return ""

def convert_citations(text):
    """APA (Author, Year) -> Springer (Author Year); '&' -> 'and'."""
    t = text.replace(" & ", " and ")
    # remove the comma before a 4-digit year (optional letter suffix), but only
    # when the char before the comma is NOT a digit (protects 1,322 etc.)
    t = re.sub(r"(?<!\d),\s+(\d{4}[a-z]?)\b", r" \1", t)
    return t

def add_rich_runs(paragraph, text, bold=False, convert=False):
    """Add runs to paragraph, honouring [[I]]..[[/I]] italic sentinels."""
    if convert:
        text = convert_citations(text)
    parts = re.split(r"(\[\[I\]\]|\[\[/I\]\])", text)
    ital = False
    for part in parts:
        if part == "[[I]]":
            ital = True
            continue
        if part == "[[/I]]":
            ital = False
            continue
        if part == "":
            continue
        run = paragraph.add_run(part)
        if ital:
            run.italic = True
        if bold:
            run.bold = True

def new_para(doc, style=None):
    """Create a paragraph at end of body (before sectPr); return Paragraph.
    'Normal'/None leaves the default body style (not always enumerated)."""
    p = doc.add_paragraph()
    if style and style != "Normal" and style in STYLE_BY_NAME:
        p.style = STYLE_BY_NAME[style]
    return p

def move_before(p, anchor_el):
    anchor_el.addprevious(p._p)

def set_heading_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for extra in runs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        p.add_run(new_text)

# ---------------------------------------------------------------------------
# 0. Copy the Methods/Results docx -> working master (preserves images/tables)
# ---------------------------------------------------------------------------
shutil.copyfile(SRC_MR, TMP)
doc = Document(TMP)
body = doc.element.body

# Robust name->style map (avoids python-docx get_by_name casing quirks).
STYLE_BY_NAME = {}
for s in doc.styles:
    try:
        STYLE_BY_NAME[s.name] = s
    except Exception:
        pass

def get_style(name):
    return STYLE_BY_NAME[name]

# ---------------------------------------------------------------------------
# 1. Locate the anchor "3  The Resource: Corpus Construction" and strip
#    everything before it (title block, scope, Methods divider + bridge).
# ---------------------------------------------------------------------------
children = list(body)
anchor = None
for ch in children:
    if ch.tag == qn('w:p'):
        t = ptext(ch).strip()
        if t.startswith("3") and "Resource" in t and "Corpus" in t:
            anchor = ch
            break
if anchor is None:
    raise SystemExit("ERROR: could not find Section 3 anchor.")

for ch in children:
    if ch is anchor:
        break
    body.remove(ch)

# ---------------------------------------------------------------------------
# 2. Strip the "References (methods and tools)" section (heading + following
#    siblings, up to but excluding the trailing sectPr).
# ---------------------------------------------------------------------------
children = list(body)
ref_start = None
for ch in children:
    if ch.tag == qn('w:p') and ptext(ch).strip().lower().startswith("references (methods"):
        ref_start = ch
        break
if ref_start is not None:
    started = False
    for ch in list(body):
        if ch is ref_start:
            started = True
        if started:
            if ch.tag == qn('w:sectPr'):
                continue
            body.remove(ch)

# ---------------------------------------------------------------------------
# 3. Normalise heading levels across the kept Methods/Results region.
#    Methods part (before "Results (Section 6)"): H2->H1 (sections 3/4/5),
#    H3->H2 (their subsections). Results part: rename divider to "6  Results"
#    (keep H1); keep 6.x at H2 and 6.2.1 at H3.
# ---------------------------------------------------------------------------
in_results = False
for ch in list(body):
    if ch.tag != qn('w:p'):
        continue
    p = Paragraph(ch, doc)
    sn = style_name(p)
    t = ptext(ch).strip()
    if sn == "Heading 1" and t.lower().startswith("results"):
        set_heading_text(p, "6  Results")
        in_results = True
        continue
    if not in_results:
        if sn == "Heading 2":
            # promote top-level numbered sections (3,4,5 -> no dot in number)
            if re.match(r"^\d+\s", t) and not re.match(r"^\d+\.", t):
                p.style = get_style("Heading 1")
        elif sn == "Heading 3":
            p.style = get_style("Heading 2")
    # in_results: leave 6.x at H2 and 6.2.1 at H3 unchanged

# ---------------------------------------------------------------------------
# 3b. Insert the corpus-provenance paragraph as a 2nd paragraph of Section 3.1.
# ---------------------------------------------------------------------------
prov_anchor = None
for ch in list(body):
    if ch.tag == qn('w:p') and ptext(ch).startswith("The human side of the corpus consists of 119"):
        prov_anchor = ch
        break
if prov_anchor is None:
    print("WARN: 3.1 paragraph not found for provenance insertion")
else:
    pp = new_para(doc, "Normal")
    add_rich_runs(pp, C.PROVENANCE_31)
    prov_anchor.addnext(pp._p)

# ---------------------------------------------------------------------------
# 4. Build front matter + Section 1 + Section 2, inserted before the anchor.
# ---------------------------------------------------------------------------
# 4a. Title
p = new_para(doc, "Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_rich_runs(p, C.TITLE)
move_before(p, anchor)

# 4b. Author
p = new_para(doc, "Normal")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(C.AUTHOR)
r.bold = True
r.font.size = Pt(13)
move_before(p, anchor)

# 4c. Affiliation lines (placeholders)
for line in C.AFFIL_LINES:
    p = new_para(doc, "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    move_before(p, anchor)

# 4d. Abstract
p = new_para(doc, "Heading 1")
set_heading_text(p, "Abstract")
move_before(p, anchor)
p = new_para(doc, "Normal")
add_rich_runs(p, C.ABSTRACT)
move_before(p, anchor)

# 4e. Keywords
p = new_para(doc, "Normal")
kr = p.add_run("Keywords  ")
kr.bold = True
p.add_run("· ".join(k + " " for k in C.KEYWORDS).strip())
move_before(p, anchor)

# 4f. Page break before the body
p = new_para(doc, "Normal")
p.add_run().add_break(WD_BREAK.PAGE)
move_before(p, anchor)

# 4g. Section 1 Introduction (new prose)
for level, text in C.INTRODUCTION:
    if level == "h1":
        p = new_para(doc, "Heading 1"); set_heading_text(p, text)
    elif level == "h2":
        p = new_para(doc, "Heading 2"); set_heading_text(p, text)
    else:
        p = new_para(doc, "Normal"); add_rich_runs(p, text)
    move_before(p, anchor)

# 4h. Section 2 Related Work (cloned from the Literature Review docx)
lr = Document(SRC_LR)
collect = False
for src_p in lr.paragraphs:
    t = src_p.text.strip()
    sn = style_name(src_p)
    if (not collect) and t.startswith("2") and "Related Work" in t:
        collect = True
    if not collect:
        continue
    if t.lower().startswith("references (literature"):
        break  # stop before the lit-review reference list
    # map heading levels: lit-review H2 ("2 Related Work") -> H1; H3 ("2.x") -> H2
    if sn == "Heading 2":
        p = new_para(doc, "Heading 1")
    elif sn == "Heading 3":
        p = new_para(doc, "Heading 2")
    else:
        p = new_para(doc, "Normal")
    # clone runs, preserving italic/bold, converting citations
    any_run = False
    for r in src_p._p.iter(qn('w:r')):
        tx = "".join(n.text or "" for n in r.iter(qn('w:t')))
        if tx == "":
            continue
        rpr = r.find(qn('w:rPr'))
        ital = rpr is not None and rpr.find(qn('w:i')) is not None
        bold = rpr is not None and rpr.find(qn('w:b')) is not None
        run = p.add_run(convert_citations(tx))
        if ital:
            run.italic = True
        if bold:
            run.bold = True
        any_run = True
    if not any_run and t:
        p.add_run(convert_citations(t))
    move_before(p, anchor)

# ---------------------------------------------------------------------------
# 5. Append back matter: Discussion, Conclusion, Declarations, References,
#    Appendix A.  (These append at the end, before the trailing sectPr.)
# ---------------------------------------------------------------------------
def append_blocks(blocks):
    for level, text in blocks:
        if level == "h1":
            p = new_para(doc, "Heading 1"); set_heading_text(p, text)
        elif level == "h2":
            p = new_para(doc, "Heading 2"); set_heading_text(p, text)
        else:
            p = new_para(doc, "Normal"); add_rich_runs(p, text)

append_blocks(C.DISCUSSION)
append_blocks(C.CONCLUSION)
append_blocks(C.DECLARATIONS)

# References
p = new_para(doc, "Heading 1"); set_heading_text(p, "References")
for ref in C.REFERENCES:
    p = new_para(doc, "Normal")
    pf = p.paragraph_format
    pf.left_indent = Pt(18)
    pf.first_line_indent = Pt(-18)  # hanging indent
    pf.space_after = Pt(4)
    p.add_run(ref)

# Appendix A
append_blocks(C.APPENDIX_INTRO)

def add_prompt(title, path):
    p = new_para(doc, "Heading 2"); set_heading_text(p, title)
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    for line in lines:
        p = new_para(doc, "Normal")
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)

add_prompt("A.1  Restricted (close-rewrite) prompt", os.path.join(ROOT, "prompts", "restricted.txt"))
add_prompt("A.2  Free prompt", os.path.join(ROOT, "prompts", "free.txt"))

# ---------------------------------------------------------------------------
# 6. Continuous line numbers (helpful for reviewers).
# ---------------------------------------------------------------------------
sectPr = body.find(qn('w:sectPr'))
if sectPr is not None and sectPr.find(qn('w:lnNumType')) is None:
    ln = OxmlElement('w:lnNumType')
    ln.set(qn('w:countBy'), '1')
    ln.set(qn('w:restart'), 'continuous')
    ln.set(qn('w:distance'), '360')
    ref = sectPr.find(qn('w:pgMar'))
    if ref is None:
        ref = sectPr.find(qn('w:pgSz'))
    if ref is not None:
        ref.addnext(ln)
    else:
        sectPr.append(ln)

# ---------------------------------------------------------------------------
# 6b. Centred page number in the footer.
# ---------------------------------------------------------------------------
def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(instr); run._r.append(f2)

try:
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_number(fp)
except Exception as e:
    print("WARN footer:", e)

# ---------------------------------------------------------------------------
# 7. Save
# ---------------------------------------------------------------------------
doc.save(OUT)
try:
    os.remove(TMP)
except OSError:
    pass

# quick abstract word count
abw = len(re.sub(r"\[\[/?I\]\]", "", C.ABSTRACT).split())
print("Saved:", OUT)
print("Abstract word count:", abw, "(target 150-250)")
print("References:", len(C.REFERENCES))
